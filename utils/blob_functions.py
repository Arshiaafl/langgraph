from azure.storage.blob import BlobServiceClient
import os
from dotenv import load_dotenv
from io import BytesIO
from docx import Document

# Load environment variables it is important to do this before importing any other modules
load_dotenv()
connection_string = os.getenv("AZURE_STORAGE_CONNECTION_STRING")

if not connection_string:
    print("Error: AZURE_STORAGE_CONNECTION_STRING environment variable not set")
    exit(1)

def connect_to_blob_storage():
    try:
        # Create the BlobServiceClient
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        # Define container
        container_name = "verify-dev"
        container_client = blob_service_client.get_container_client(container_name)
        print("Connected to blob storage") 
        return container_client
            
    except Exception as e:
        print("Error in connecting to blob storage")
        print(f"Error: {str(e)}")


def get_all_files_from_folder(folder_name):
    container_client = connect_to_blob_storage()
    blobs = container_client.list_blobs(name_starts_with=folder_name)
    files = []
    for blob in blobs:
        files.append(blob.name)
    return files

def get_file_content_as_text(path):
    try:
        container_client = connect_to_blob_storage()
        blob_client = container_client.get_blob_client(blob=path)
        blob_data = blob_client.download_blob()
        file_bytes = blob_data.readall()

        # Convert bytes to a file-like object
        file_stream = BytesIO(file_bytes)
        doc = Document(file_stream)

        # Extract all text from the document
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        print(f"Error retrieving file {path}: {str(e)}")
        return None
    
def get_file_content_as_docx(path):
    try:
        container_client = connect_to_blob_storage()
        blob_client = container_client.get_blob_client(blob=path)
        blob_data = blob_client.download_blob()
        file_bytes = blob_data.readall()

        # Convert bytes to a file-like object
        file_stream = BytesIO(file_bytes)
        doc = Document(file_stream)

        return doc
    except Exception as e:
        print(f"Error retrieving file {path}: {str(e)}")
        return None

def save_docx_to_azure(doc: Document, path: str):
    try:
        container_client = connect_to_blob_storage()
        blob_client = container_client.get_blob_client(blob=path)
        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)

        response = blob_client.upload_blob(stream, overwrite=True)
        print(f"✅ Saved filled document to Azure blob '{path}'")
        print(response)
        output_url = f"https://verifystore.blob.core.windows.net/verify-dev/{path}"
        return output_url
    except Exception as e:
        print(f"Error saving document to Azure: {str(e)}")
        raise e

########################################################
# S3 Equivaluent Functions
########################################################

def extract_text_from_blob_docx(blob_path: str) -> str:
    try:
        doc = get_file_content_as_docx(blob_path)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        return f"Error reading DOCX from Azure Blob: {e}"

import pdfplumber

def extract_text_from_docx(docx_path: str) -> str:
    try:
        doc = Document(docx_path)
        return "\n".join(para.text for para in doc.paragraphs)
    except Exception as e:
        return f"Error reading DOCX file: {e}"

# Function to extract text from PDF file (local)
def extract_text_from_pdf(pdf_path: str) -> str:
    try:
        with pdfplumber.open(pdf_path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception as e:
        return f"Error reading PDF file: {e}"

def extract_text_from_blob_pdf(blob_path: str) -> str:
    try:
        container_client = connect_to_blob_storage()
        blob_client = container_client.get_blob_client(blob=blob_path)
        blob_data = blob_client.download_blob()
        file_bytes = blob_data.readall()

        buffer = BytesIO(file_bytes)
        with pdfplumber.open(buffer) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception as e:
        return f"Error reading PDF from Azure Blob: {e}"

def save_text_to_blob_as_docx(text: str, filename: str) -> str:
    try:
        doc = Document()
        doc.add_paragraph(text)
        return save_docx_to_azure(doc, filename)
    except Exception as e:
        print(f"Error saving DOCX to Azure: {e}")
        raise

from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet

def save_text_to_blob_as_pdf(text: str, filename: str) -> str:
    try:
        container_client = connect_to_blob_storage()
        blob_client = container_client.get_blob_client(blob=filename)

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        flowables = [Paragraph(text.replace('\n', '<br />'), styles['Normal'])]
        doc.build(flowables)

        buffer.seek(0)
        blob_client.upload_blob(buffer, overwrite=True)

        output_url = f"https://verifystore.blob.core.windows.net/verify-dev/{filename}"
        return output_url
    except Exception as e:
        print(f"Error saving PDF to Azure: {e}")
        raise

def save_text_output_to_blob(text: str, docx_filename: str, pdf_filename: str) -> tuple:
    docx_url = save_text_to_blob_as_docx(text, docx_filename)
    pdf_url = save_text_to_blob_as_pdf(text, pdf_filename)
    return docx_url, pdf_url